"""File upload processing -- extract text from PDF, DOCX, plain-text, and images.

Matches the n8n Switch4 node which routes files by MIME type:
  PDF        -> pdfplumber (text + tables)
  Plain Text -> direct read
  Image      -> Pillow metadata (OCR deferred)
  DOCX       -> python-docx (paragraphs + tables)
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Literal

import pdfplumber
from docx import Document as DocxDocument
from PIL import Image

# Threshold (chars) below which extraction is considered "partial"
_PARTIAL_THRESHOLD = 200


@dataclass(frozen=True)
class ExtractedFile:
    filename: str
    mime_type: str
    extracted_text: str
    char_count: int
    has_tables: bool = False
    extraction_quality: Literal["full", "partial", "metadata_only", "empty"] = "full"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _table_to_markdown(rows: list[list[str | None]]) -> str:
    """Convert a 2D table (from pdfplumber or python-docx) to Markdown.

    Handles: empty rows, inconsistent column counts, and pipe characters
    in cell content.
    """
    if not rows:
        return ""
    # Filter completely empty rows, escape pipes, normalize to strings
    clean = [
        [(cell or "").replace("|", "\\|") for cell in row]
        for row in rows
        if any(cell for cell in row)
    ]
    if not clean:
        return ""
    # Normalize column count across all rows
    max_cols = max(len(row) for row in clean)
    clean = [row + [""] * (max_cols - len(row)) for row in clean]

    header = "| " + " | ".join(clean[0]) + " |"
    sep = "| " + " | ".join("---" for _ in clean[0]) + " |"
    body_lines = ["| " + " | ".join(row) + " |" for row in clean[1:]]
    return "\n".join([header, sep] + body_lines)


def _quality(text: str) -> Literal["full", "partial", "empty"]:
    """Score extraction quality for text-based extractors.

    Note: ``"metadata_only"`` is assigned directly by the image extractor,
    not through this function.
    """
    chars = len(text)
    if chars == 0:
        return "empty"
    if chars < _PARTIAL_THRESHOLD:
        return "partial"
    return "full"


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------

def extract_pdf(file_bytes: bytes, filename: str) -> ExtractedFile:
    parts: list[str] = []
    table_count = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()

            # Extract text excluding table regions to avoid duplication
            if tables:
                table_bboxes = [t.bbox for t in page.find_tables()]
                filtered_page = page
                for bbox in table_bboxes:
                    filtered_page = filtered_page.outside_bbox(bbox)
                page_text = (filtered_page.extract_text() or "").strip()
            else:
                page_text = (page.extract_text() or "").strip()

            if page_text:
                parts.append(page_text)

            for table in tables:
                md = _table_to_markdown(table)
                if md:
                    parts.append(f"[Table]\n{md}")
                    table_count += 1

    text = "\n\n".join(parts).strip()
    return ExtractedFile(
        filename=filename,
        mime_type="application/pdf",
        extracted_text=text,
        char_count=len(text),
        has_tables=table_count > 0,
        extraction_quality=_quality(text),
    )


def extract_docx(file_bytes: bytes, filename: str) -> ExtractedFile:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts: list[str] = []
    table_count = 0

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            seen: set[int] = set()
            row_data: list[str] = []
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen:
                    row_data.append("")  # merged cell placeholder
                else:
                    seen.add(cell_id)
                    row_data.append(cell.text.strip())
            rows.append(row_data)
        md = _table_to_markdown(rows)
        if md:
            parts.append(f"[Table]\n{md}")
            table_count += 1

    text = "\n\n".join(parts).strip()
    return ExtractedFile(
        filename=filename,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        extracted_text=text,
        char_count=len(text),
        has_tables=table_count > 0,
        extraction_quality=_quality(text),
    )


def extract_plain_text(file_bytes: bytes, filename: str) -> ExtractedFile:
    text = file_bytes.decode("utf-8", errors="replace").strip()
    return ExtractedFile(
        filename=filename,
        mime_type="text/plain",
        extracted_text=text,
        char_count=len(text),
        extraction_quality=_quality(text),
    )


def extract_image_description(file_bytes: bytes, filename: str) -> ExtractedFile:
    """Return basic image metadata. Vision/OCR analysis is deferred to the LLM."""
    with Image.open(io.BytesIO(file_bytes)) as img:
        description = f"Image: {filename} ({img.format}, {img.size[0]}x{img.size[1]})"
        mime = f"image/{(img.format or 'unknown').lower()}"
    return ExtractedFile(
        filename=filename,
        mime_type=mime,
        extracted_text=description,
        char_count=len(description),
        extraction_quality="metadata_only",
    )


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_MIME_HANDLERS = {
    "application/pdf": extract_pdf,
    "text/plain": extract_plain_text,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_docx,
    "application/zip": extract_docx,  # n8n treats DOCX uploads as zip
}


def process_file(file_bytes: bytes, filename: str, mime_type: str) -> ExtractedFile:
    """Route *file_bytes* to the correct extractor based on *mime_type*."""

    if mime_type in _MIME_HANDLERS:
        return _MIME_HANDLERS[mime_type](file_bytes, filename)

    if mime_type.startswith("image/"):
        return extract_image_description(file_bytes, filename)

    # Fallback: treat as plain text
    return extract_plain_text(file_bytes, filename)
