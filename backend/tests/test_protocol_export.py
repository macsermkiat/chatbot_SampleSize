"""Tests for protocol export service (DOCX generation)."""

from __future__ import annotations

import io
from unittest.mock import patch

import pytest
from docx import Document

from app.services.protocol_export import (
    _build_protocol_sections,
    _parse_summary_sections,
    generate_docx,
    generate_pdf,
    generate_protocol,
)

SESSION_ID = "550e8400-e29b-41d4-a716-446655440000"

# A realistic structured summary as the LLM would produce
STRUCTURED_SUMMARY = """\
## Background and Rationale
Statins are widely prescribed. There is a gap in long-term cohort data [1].

## Research Question
Does long-term statin use reduce dementia risk in adults over 65?

## Study Design and Methodology
A retrospective cohort design using Target Trial Emulation is recommended.

## Statistical Analysis Plan
Power analysis: n=1,200 per arm for 80% power at alpha=0.05.

## Unresolved Issues and Recommendations
Duration of follow-up needs clarification.

## References
1. Smith et al. Statins and Dementia. Lancet. 2024.

## Executive Summary (1-page)
This consultation explored statin-dementia risk using a cohort design."""

SAMPLE_MESSAGES = [
    {"role": "user", "content": "What study design for statin dementia?", "phase": ""},
    {
        "role": "assistant",
        "content": "Based on literature review, there are gaps in long-term cohort studies.",
        "phase": "research_gap",
    },
    {
        "role": "assistant",
        "content": "A retrospective cohort design using Target Trial Emulation is recommended.",
        "phase": "methodology",
    },
    {
        "role": "assistant",
        "content": "Power analysis: n=1,200 per arm for 80% power at alpha=0.05.",
        "phase": "biostatistics",
    },
]

MESSAGES_WITH_CITATIONS = [
    {"role": "user", "content": "Analyze this topic", "phase": ""},
    {
        "role": "assistant",
        "content": "See [Study X](https://example.com/study-x) for background on CONSORT.",
        "phase": "research_gap",
    },
    {
        "role": "assistant",
        "content": "Based on [Protocol Y](https://example.com/protocol-y).",
        "phase": "methodology",
    },
]


class TestParseSummarySections:
    def test_parses_markdown_headings(self):
        sections = _parse_summary_sections(STRUCTURED_SUMMARY)
        headings = [s["heading"] for s in sections]
        assert "Background and Rationale" in headings
        assert "Research Question" in headings
        assert "Statistical Analysis Plan" in headings

    def test_plain_text_returns_single_overview(self):
        sections = _parse_summary_sections("Just plain text, no headings.")
        assert len(sections) == 1
        assert sections[0]["heading"] == "Overview"

    def test_empty_string_returns_empty(self):
        sections = _parse_summary_sections("")
        assert len(sections) == 0


class TestBuildProtocolSections:
    def test_structured_summary_produces_multiple_sections(self):
        sections = _build_protocol_sections(STRUCTURED_SUMMARY, [], SESSION_ID)
        headings = [s["heading"] for s in sections]
        assert "Background and Rationale" in headings
        assert "Study Design and Methodology" in headings
        assert "Statistical Analysis Plan" in headings

    def test_plain_summary_falls_back_to_single_section(self):
        sections = _build_protocol_sections("Just the summary", [], SESSION_ID)
        assert len(sections) == 1
        assert sections[0]["content"] == "Just the summary"

    def test_no_duplicate_references_if_summary_has_one(self):
        sections = _build_protocol_sections(STRUCTURED_SUMMARY, MESSAGES_WITH_CITATIONS, SESSION_ID)
        ref_headings = [s for s in sections if s["heading"].lower() == "references"]
        assert len(ref_headings) == 1

    def test_adds_references_from_citations_when_summary_lacks_them(self):
        summary_no_refs = "## Background\nSome background.\n\n## Methods\nSome methods."
        sections = _build_protocol_sections(summary_no_refs, MESSAGES_WITH_CITATIONS, SESSION_ID)
        headings = [s["heading"] for s in sections]
        assert "References" in headings


class TestGeneratePdf:
    def test_returns_valid_pdf_bytes(self):
        result = generate_pdf(STRUCTURED_SUMMARY, SAMPLE_MESSAGES, SESSION_ID)
        assert isinstance(result, bytes)
        assert len(result) > 0
        assert result[:5] == b"%PDF-"

    def test_pdf_with_empty_messages(self):
        result = generate_pdf("Just the summary", [], SESSION_ID)
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    def test_pdf_with_unicode_content(self):
        messages = [
            {
                "role": "assistant",
                "content": "Effect size: d = 0.5, CI: 0.3\u20130.7, p < 0.001",
                "phase": "biostatistics",
            },
        ]
        result = generate_pdf("Summary with \u2014 em-dash", messages, SESSION_ID)
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"


class TestGenerateDocx:
    def test_returns_valid_docx_bytes(self):
        result = generate_docx(STRUCTURED_SUMMARY, SAMPLE_MESSAGES, SESSION_ID)
        assert isinstance(result, bytes)
        assert len(result) > 0
        # Verify it's a valid DOCX (ZIP format, starts with PK)
        assert result[:2] == b"PK"

    def test_docx_contains_title(self):
        result = generate_docx(STRUCTURED_SUMMARY, [], SESSION_ID)
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        full_text = "\n".join(texts)
        assert "Research Protocol" in full_text

    def test_docx_contains_structured_sections(self):
        result = generate_docx(STRUCTURED_SUMMARY, [], SESSION_ID)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Background and Rationale" in full_text
        assert "Statistical Analysis Plan" in full_text
        assert "Statins are widely prescribed" in full_text

    def test_docx_contains_disclaimer(self):
        result = generate_docx(STRUCTURED_SUMMARY, [], SESSION_ID)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Disclaimer" in full_text
        assert "verified" in full_text.lower()

    def test_docx_contains_session_id_prefix(self):
        result = generate_docx(STRUCTURED_SUMMARY, [], SESSION_ID)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert SESSION_ID[:8] in full_text


class TestGenerateProtocol:
    def test_docx_format(self):
        data, content_type, filename = generate_protocol(
            "Summary", SAMPLE_MESSAGES, SESSION_ID, format="docx"
        )
        assert content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert filename.endswith(".docx")
        assert SESSION_ID[:8] in filename
        assert isinstance(data, bytes)

    def test_pdf_format_returns_bytes(self):
        data, content_type, filename = generate_protocol(
            "Summary", SAMPLE_MESSAGES, SESSION_ID, format="pdf"
        )
        assert content_type == "application/pdf"
        assert filename.endswith(".pdf")
        assert SESSION_ID[:8] in filename
        assert isinstance(data, bytes)
        assert len(data) > 0
        # PDF magic bytes
        assert data[:5] == b"%PDF-"

    def test_default_format_is_docx(self):
        _, content_type, filename = generate_protocol("Summary", [], SESSION_ID)
        assert filename.endswith(".docx")


class TestReferencesInExport:
    def test_docx_contains_references_section_when_links_present(self):
        summary_no_refs = "## Background\nSome background."
        result = generate_docx(summary_no_refs, MESSAGES_WITH_CITATIONS, SESSION_ID)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "References" in full_text
        assert "Study X" in full_text
        assert "example.com/study-x" in full_text

    def test_no_references_section_when_no_citations(self):
        plain_messages = [
            {"role": "assistant", "content": "No links here.", "phase": "methodology"},
        ]
        result = generate_docx("Just a summary", plain_messages, SESSION_ID)
        doc = Document(io.BytesIO(result))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "References" not in headings

    def test_references_appear_before_disclaimer(self):
        summary_no_refs = "## Background\nSome background."
        result = generate_docx(summary_no_refs, MESSAGES_WITH_CITATIONS, SESSION_ID)
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Find positions
        ref_idx = next(i for i, t in enumerate(texts) if "References" in t)
        disc_idx = next(i for i, t in enumerate(texts) if "Disclaimer" in t)
        assert ref_idx < disc_idx

    def test_static_refs_included_for_keyword_match(self):
        summary_no_refs = "## Background\nSome background."
        result = generate_docx(summary_no_refs, MESSAGES_WITH_CITATIONS, SESSION_ID)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # CONSORT keyword in messages should trigger static ref
        assert "CONSORT" in full_text
